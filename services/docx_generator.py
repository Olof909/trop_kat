from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
        el = OxmlElement(tag)
        if tag == 'w:fldChar':
            el.set(qn('w:fldCharType'), 'begin' if text is None and not run._r.findall(qn('w:fldChar')) else 'end')
        elif tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        run._r.append(el)


def _page_number_footer(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    for fld_type in ('begin', 'end'):
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), fld_type)
        if fld_type == 'begin':
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._r.append(fldChar)
            run._r.append(instrText)
        else:
            run._r.append(fldChar)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def generate_docx(articles: list, document_title: str) -> bytes:
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)

    # Default style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Page numbers in footer
    _page_number_footer(section)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(32)
    run = title_para.add_run(document_title.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # ── Articles ───────────────────────────────────────────────────────────────
    for article in articles:
        article_id = article.id if hasattr(article, 'id') else article['id']
        title = article.title if hasattr(article, 'title') else article['title']
        content = article.content if hasattr(article, 'content') else article['content']

        # Build header text
        if article_id == 0:
            header_text = title.strip()
        else:
            header_text = f'Άρθρο {article_id}'
            if title.strip() and title.strip() != header_text:
                header_text += f' – {title.strip()}'

        # Article header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_before = Pt(24)
        header_para.paragraph_format.space_after = Pt(8)
        run = header_para.add_run(header_text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Article body — each non-empty line becomes a justified paragraph
        lines = [ln for ln in content.split('\n')]
        for line in lines:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = Pt(20)
            para.paragraph_format.first_line_indent = Cm(1.25)
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
